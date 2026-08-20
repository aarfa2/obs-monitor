# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).with_name("OBS机群监控系统-项目汇报.docx")

BLUE = RGBColor(0x1E, 0x40, 0xAF)
NAVY = RGBColor(0x1E, 0x3A, 0x8A)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x64, 0x74, 0x8B)
TH_BG = "EFF6FF"
ALT_BG = "FAFAFA"
META_BG = "F8FAFC"


def set_run_font(run, name="微软雅黑", size=11, bold=False, color=DARK, east_asia=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    ea = east_asia or name
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), ea)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CBD5E1")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, *, bold=False, header=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold or header, color=NAVY if header else DARK)
    shade_cell(cell, TH_BG if header else "FFFFFF")
    set_cell_borders(cell)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, header=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            set_cell_text(table.rows[r_i + 1].cells[c_i], val, bold=(c_i == 0 and len(headers) == 2))
            if r_i % 2 == 1:
                shade_cell(table.rows[r_i + 1].cells[c_i], ALT_BG)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 0:
            set_run_font(run, size=22, bold=True, color=RGBColor(0x11, 0x11, 0x11))
        elif level == 1:
            set_run_font(run, size=14, bold=True, color=BLUE)
        else:
            set_run_font(run, size=12, bold=True, color=DARK)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    return p


def para(doc, parts, *, space_after=8):
    """parts: str or list of (text, bold)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if isinstance(parts, str):
        parts = [(parts, False)]
    for text, bold in parts:
        run = p.add_run(text)
        set_run_font(run, size=11, bold=bold)
    return p


def add_list(doc, items, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.45
        if isinstance(item, str):
            run = p.add_run(item)
            set_run_font(run, size=11)
        else:
            for text, bold in item:
                run = p.add_run(text)
                set_run_font(run, size=11, bold=bold)


def quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=RGBColor(0x33, 0x41, 0x55))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2563EB")
    pBdr.append(left)
    pPr.append(pBdr)


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    set_run_font(run, name="Consolas", east_asia="微软雅黑", size=9, color=DARK)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F1F5F9")
    shd.set(qn("w:val"), "clear")
    p._p.get_or_add_pPr().append(shd)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "微软雅黑")
    rFonts.set(qn("w:hAnsi"), "微软雅黑")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("OBS 机群监控系统 — 项目汇报")
    set_run_font(run, size=22, bold=True, color=RGBColor(0x11, 0x11, 0x11))
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2563EB")
    pBdr.append(bottom)
    title._p.get_or_add_pPr().append(pBdr)

    meta = [
        ("汇报人：", "Scott"),
        ("汇报日期：", "2026 年 8 月 20 日"),
        ("项目版本：", "v0.1.0（局域网场景已可实际使用）"),
        ("适用范围：", "直播/录播现场多台 OBS 推流电脑的集中监控"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(1)
        r1 = p.add_run(label)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), META_BG)
        shd.set(qn("w:val"), "clear")
        p._p.get_or_add_pPr().append(shd)
    doc.add_paragraph()

    heading(doc, "一、摘要")
    para(doc, [
        ("本项目是一套面向", False),
        ("生产推流机房", True),
        ("的值班监控系统，解决「十几台 OBS 电脑分散运行、故障发现慢、排障证据散」的问题。", False),
    ])
    para(doc, [
        ("系统采用", False),
        ("中心看板 + 各机轻量采集器", True),
        ("架构：值班人员在一个网页上即可查看全部机器是否在线、是否在推流、是否存在渲染/编码/网络压力，并在异常时通过 Webhook 自动通知值班群。", False),
        ("只监控、不控制", True),
        ("——不能通过本系统修改 OBS 配置或停推流，避免误操作风险。核心功能已开发完成，可在局域网内部署试用。", False),
    ])

    heading(doc, "二、背景与痛点")
    para(doc, "现场典型情况：多台电脑各自用 OBS 推流到服务器，采集卡、摄像头、编码器、上行网络任一环节出问题，画面就会卡、花、停，甚至整机静默。")
    add_table(doc, ["现状问题", "影响"], [
        ["需逐台远程桌面或到机房查看", "机器一多，人力跟不上"],
        ["往往观众先投诉才发现", "故障发现滞后，影响播出质量"],
        ["出事后翻各机 OBS 日志", "时间对不齐、容易漏、排障慢"],
        ["若开放 OBS 远程控制接口", "存在误操作停推流等生产风险"],
    ])
    para(doc, [
        ("核心需求：", True),
        ("看得见（集中看板）→ 来得及（自动报警）→ 查得清（日志检索与掉帧分类）→ 控得住风险（监控与控制分离）。", False),
    ])

    heading(doc, "三、建设目标")
    add_list(doc, [
        [("集中可视：", True), ("全部 OBS 机器推流状态、健康指标一屏呈现", False)],
        [("主动告警：", True), ("停推、重连、失联等异常自动通知，不依赖人盯屏", False)],
        [("快速定位：", True), ("区分渲染掉帧、编码掉帧、网络丢帧，支持近 24 小时日志检索", False)],
        [("安全边界：", True), ("只读采集，不向 OBS 下发任何控制指令", False)],
    ], ordered=True)

    heading(doc, "四、系统能力与业务价值")
    heading(doc, "4.1 已具备功能", level=2)
    add_table(doc, ["模块", "能力说明"], [
        ["机群总览", "每台机器一张卡片：在线/离线、推流状态、重连、压力等级、码率、CPU；异常机器自动置顶"],
        ["单机详情", "场景与源列表、码率/帧率/拥塞、三类掉帧率、录像状态、历史曲线"],
        ["智能报警", "OBS 失联、意外停推、持续重连、采集器心跳超时；支持 Webhook（可接企业微信/钉钉）；恢复后自动发恢复通知"],
        ["日志中心", "汇聚各机 OBS 原生日志，按分类/级别/关键字检索；敏感信息脱敏；中心保留约 24 小时"],
    ])

    heading(doc, "4.2 三类掉帧（排障核心指标）", level=2)
    para(doc, "比单一「GPU 占用」更有现场指导意义：")
    add_list(doc, [
        [("渲染掉帧", True), (" → 场景过重、浏览器源、滤镜等合成压力", False)],
        [("编码掉帧", True), (" → 编码器过载", False)],
        [("网络丢帧", True), (" → 上行带宽不足或服务器侧不稳", False)],
    ])
    quote(doc, "OBS 官方接口不提供 GPU 占用数据，系统不编造该指标；渲染掉帧是判断「显卡/合成压力」的可靠依据。")

    heading(doc, "4.3 业务价值", level=2)
    add_table(doc, ["价值维度", "预期效果"], [
        ["故障发现", "从「观众反馈/人工巡机」→「看板变红 + 即时通知」"],
        ["排障效率", "先看掉帧类型和日志分类，减少盲目重启和远程桌面"],
        ["规模扩展", "10 台与 20 台同一套中心，按「一机一采集器」线性扩展"],
        ["生产安全", "监控面与控制面分离，OBS 配置权仍在现场机"],
    ])

    heading(doc, "五、技术方案概要")
    code_block(
        doc,
        "OBS 电脑 × N              监控服务器（Hub）              值班终端\n"
        "┌──────────────┐         ┌─────────────────┐         ┌──────────┐\n"
        "│ OBS 推流      │         │ 网页看板         │         │ 浏览器    │\n"
        "│ 采集器 Agent  │ ──WS──► │ 日志存储 + 报警  │ ◄─HTTP─ │ 机群→单机 │\n"
        "│ 只连本机 OBS  │         │ 端口 8787        │         └──────────┘\n"
        "└──────────────┘         └─────────────────┘",
    )
    para(doc, [("设计要点：", True)])
    add_list(doc, [
        [("Agent（采集器）：", True), ("跑在每台 OBS 电脑，无界面、轻量（约数十 MB 内存），每秒约 1 次只读查询，不装 Electron，避免与推流抢资源", False)],
        [("Hub（中心）：", True), ("跑在监控服务器一台，对外提供唯一看板地址，接收全部采集器数据，存日志、发 Webhook", False)],
        [("技术栈：", True), ("TypeScript + Fastify（后端）+ React（看板），基于 OBS 28+ 官方 WebSocket 接口", False)],
    ])
    para(doc, [
        ("安全设计：", True),
        ("各机 OBS WebSocket 仅监听本机（127.0.0.1:4455），采集器只读采集后出站上报；中心", False),
        ("不会", True),
        ("向采集器下发控制命令。", False),
    ])

    heading(doc, "六、当前进展")
    add_table(doc, ["项", "状态"], [
        ["机群看板 / 单机详情页", "已完成"],
        ["三类掉帧与压力判定", "已完成"],
        ["Webhook 报警（含测试按钮）", "已完成"],
        ["24 小时日志汇聚与分类检索", "已完成"],
        ["采集器 / 中心分离部署", "已完成"],
        ["开机服务 / 安装包", "待做"],
        ["对接值班群 Webhook", "待配置"],
        ["外网远程查看", "按需评估（仍不建议暴露 OBS 端口）"],
    ])
    para(doc, [
        ("项目成熟度：", True),
        ("v0.1.0，局域网场景已可实际部署使用；尚未做大规模铺机与运维流程固化。", False),
    ])

    heading(doc, "七、部署与运维")
    heading(doc, "7.1 部署方式", level=2)
    add_table(doc, ["角色", "部署位置", "启动方式"], [
        ["中心 Hub", "监控服务器 1 台", "npm run build && npm start，看板 http://服务器IP:8787"],
        ["采集器 Agent", "每台 OBS 电脑 1 个", "npm run start:agent，配置中心地址、Token、显示名、本机 OBS 密码"],
    ])
    heading(doc, "7.2 运维建议", level=2)
    add_list(doc, [
        [("各机 OBS WebSocket ", False), ("不对公网/全域网开放", True), ("，仅本机访问", False)],
        [("采集器与中心共用 Token；每台机器自动生成唯一 ID，", False), ("禁止复制", True), ("到其他机器", False)],
        [("日志默认保留 24 小时，中心上限约 20 万条，避免磁盘堆满", False)],
        [("采集器可配置为 Windows 开机服务（NSSM / 任务计划程序）", False)],
    ], ordered=True)

    heading(doc, "八、风险与系统边界")
    heading(doc, "8.1 明确能做到的", level=2)
    add_list(doc, [
        "实时监控推流状态、掉帧、源列表、OBS 日志",
        "采集器/整机失联检测（含电脑死机、进程被杀）",
        "局域网内十余台机器规模完全可承受",
    ])
    heading(doc, "8.2 明确做不到的（需管理预期）", level=2)
    add_table(doc, ["边界", "说明"], [
        ["观众端 CDN 卡顿", "本系统监控的是推流端，不是 CDN/播放端"],
        ["每路摄像机画面监看", "不是多路视频监看系统，是 OBS 运行状态监控"],
        ["推流服务器自身状态", "不能替代 RTMP/SRT 服务器侧监控"],
        ["GPU 占用百分比", "OBS 接口不提供，系统不编造"],
        ["远程控制 OBS", "刻意不做，避免生产事故"],
    ])
    heading(doc, "8.3 资源影响评估", level=2)
    para(doc, [
        ("采集器为无界面进程，约每秒 1 次轻量只读查询，", False),
        ("正常情况下对 OBS 编码影响可忽略", True),
        ("。", False),
    ])

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(16)
    run = foot.add_run("本报告基于 OBS Monitor 项目整理 · 2026-08-20")
    set_run_font(run, size=9, color=GRAY)

    doc.save(OUT)
    print(f"saved {OUT}")


if __name__ == "__main__":
    main()
